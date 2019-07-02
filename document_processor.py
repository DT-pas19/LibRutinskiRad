# @see https://docxtpl.readthedocs.io/en/latest/#introduction

import os
import re
import shutil
from datetime import datetime
from itertools import groupby
from operator import itemgetter
from pathlib import Path
from typing import List, Dict, Tuple

from PyQt5.QtWidgets import QMessageBox
from docx import Document
from pip._internal.utils.misc import file_contents

from document_structure import Work
from helpers import remove_keywords, get_bracket_less_line, replace_whitespace, insert_title_page
from title_page_creation_kit import make_title_page

key_words = {'group': 'групп', 'faculty': 'факультет', 'chair': 'кафедра', 'topic': ('на тему', 'тема'), 'student': ('выполнил', 'студент'), 'qual': ('квалификация', 'степень'), 'profile': 'профиль', 'variant': 'вариант', 'discipline': 'дисциплин', 'study': ('направлени', 'подготовки')}
numbered_filename_regexr = '\d{4,}_[бсма1]+-?[А-Я]{3,5}([зипу-]+)?\d{2,}_\d{4}_\d{1,2}_\d{1,2}'

def get_paragraph(key: str, text: List[str], include_next: Tuple[bool, int]=(False, 0)) -> str:
    """ модифицировать для tuple, получать доп поля
    Получение параграфа(-ов) с нужным ключевым словом
    :param key: ключевое слово
    :param text: текст разделенный на параграфы
    :param include_next: кортеж с информацией добавлять ли к найденному следующие параграфы
    :return: найденный текст
    """
    if key not in key_words.keys() or len(text) == 0:
        return str()
    if isinstance(key_words[key], tuple):
        result = [(index, par) for index, par in enumerate(text)
                  for k in key_words[key]
                  if k in par.lower()]
        result = list(set(result))
    else:
        result = [(index, par) for index, par in enumerate(text) if key_words[key] in par.lower()]
    final_result = '\n'.join([r for index, r in result])
    if include_next[0]:
        final_result = ''
        for ind, r in result:
            final_result += r + '\n'
            for i in range(ind + 1, ind + include_next[1] + 1):
                if i < len(text):
                    final_result += text[i] + '\n'
        final_result = final_result.strip()
    return final_result


def get_line(substr: str, text: List[str]) -> int:
    """
    Возвращает номер параграфа в тексте, в котором имеется определенная подстрока
    :param substr: подстрока
    :param text: текст
    :return: номер
    """
    if not isinstance(substr, str) or len(substr) == 0:
        return -1
    if not isinstance(text, list) or len(text) == 0:
        return -1
    line_info = [i for i, line in enumerate(text) if replace_whitespace(substr.lower()) in line.lower()]
    if len(line_info) == 0:
        line_number = -1
    else:
        line_number = next(iter(line_info))
    return line_number

# TODO filename change, manual


def get_associated_works(files: List[str], contents: List[str]):
    work_groups = []

    for document_filename in files:
        group = get_associated_work_filenames(document_filename, contents)
        work_groups.append(group)

    return work_groups


def get_associated_work_filenames(document_filename: str, contents: List[str]):
    """
    Определяет и возвращает перечень связанных работ
    Например, docx + xlsx
    Работает, если файлы имеют одни и те же наименования, или имеют одиновые префиксы до '_1', '_2', ...
    :param document_filename: имя файла документа .docx
    :param contents: содержимое папки, имена файлов
    :return:
    """
    local_document_filename = os.path.split(document_filename)[-1]
    extensionless_filename = os.path.splitext(local_document_filename)[0]
    file_group = []
    equally_named_files = [filename for filename in contents
                           if extensionless_filename == os.path.splitext(filename)[0] and not filename.endswith('.docx')]

    if re.match(numbered_filename_regexr, extensionless_filename):
        base_document_filename = '_'.join(extensionless_filename.split('_')[:-1])
        same_prefix_filenames = [filename for filename in contents
                                 if os.path.splitext(filename)[0].startswith(base_document_filename)
                                 and not filename.endswith('.docx')]
        file_group.extend(same_prefix_filenames)
    else:
        file_group.extend(equally_named_files)

    for file in file_group:
        contents.remove(file)

    file_group = list(sorted(file_group, key=lambda x: os.path.splitext(x)[-1]))
    file_group.insert(0, document_filename)

    return file_group


def read_data(folder: str, files: List[str]) -> List[Work]:
    """
    Обработка документов docx
    :param folder: исходная папка
    :param files: файлы документов
    :return: перечень работ
    """
    contents = [item for item in os.listdir(folder) if os.path.isfile(os.path.join(folder, item))]
    works = get_associated_works(files, contents)

    if len(contents) == 0:
        return

    parsed_info = []
    for documents in works:
        try:
            document = Document(os.path.join(folder, documents[0]))
        except ValueError:
            print('Ошибка чтения файла {0}'.format(documents[0]))
            continue

        first_page_text = get_document_first_page(document)
        first_page_text_str = '\n'.join(first_page_text)

        multi_value_keys = [k for key in key_words.values() if isinstance(key, tuple) for k in key]
        single_value_keys = [key for key in key_words.values() if isinstance(key, str)]
        key_words_values = multi_value_keys + single_value_keys
        contain_check = [key.lower() in first_page_text_str.lower() for key in key_words_values]

        if not any(contain_check):
            doc_info = Work(files=documents)
            parsed_info.append(doc_info)
        else:
            faculty_raw = get_paragraph('faculty', first_page_text)
            student_info_raw = get_paragraph('student', first_page_text, include_next=(True, 2))

            if len(student_info_raw) == 0:
                student_info_raw = str()
            else:
                student_info_match_ind = re.search('([А-ЯЁ][а-яё]+ ?){2,5}|(([А-ЯЁ][а-яё]+ ?) ([А-ЯЁ][. ]?){2})', student_info_raw)
                if student_info_match_ind is not None:
                    student_info_match_ind = student_info_match_ind.span()  # search_field
                    student_info_raw = student_info_raw[student_info_match_ind[0]:student_info_match_ind[1]]

            topic_raw = get_paragraph('topic', first_page_text, include_next=(True, 2))
            chair_raw = get_paragraph('chair', first_page_text, include_next=(True, 1))
            group_raw = get_paragraph('group', first_page_text)

            discipline_raw = get_paragraph('discipline', first_page_text)
            profile_raw = get_paragraph('profile', first_page_text)
            variant_raw = get_paragraph('variant', first_page_text)
            study_raw = get_paragraph('study', first_page_text)
            qual_raw = get_paragraph('qual', first_page_text)

            number_gr = [get_bracket_less_line(sub) for sub in first_page_text_str.split()
                         if get_bracket_less_line(sub).isnumeric() and len(sub) >= 6]
            if len(number_gr) > 0:
                number = next(iter(number_gr))
            else:
                file_name = os.path.split(documents[0])[1]
                number_gr = [part for part in re.split('[_ ]', file_name) if part.isnumeric() and len(part) >= 6]
                number = next(iter(number_gr)) if len(number_gr) > 0 else '0'

            faculty = remove_keywords(faculty_raw, key_words)
            chair = remove_keywords(chair_raw, key_words)
            topic = remove_keywords(topic_raw, key_words)
            group = remove_keywords(group_raw, key_words)

            if topic != str():
                topic_span = re.search('([\w\d]+ )+', topic).span()
                topic = topic[topic_span[0]:topic_span[1]].strip()

            year_match_ind = re.search('(\w{4,15} )(\d{4})', first_page_text_str)
            if year_match_ind is None:
                year_info = str()
                year = datetime.now().year
            else:
                year_match_ind = year_match_ind.span()
                year_info = first_page_text_str[year_match_ind[0]:year_match_ind[1]]
                year = re.sub('(\w{4,15} )\n?(\d{4})', '\g<2>', year_info)
                year = int(year)

            discipline = ''
            discipline_code = ''
            variant = str()
            qual = str()
            profile = str()
            study = (str(), str())
            if discipline_raw != str():
                discipline_raw = remove_keywords(discipline_raw, key_words)
                discipline_span = re.search('Б(\.?\d){3}[^\w\d]{1,5}\w+', discipline_raw)
                if discipline_span is not None:
                    discipline_span = discipline_span.span()
                    discipline_raw = discipline_raw[discipline_span[0]:discipline_span[1]]
                discipline_parts = re.split('[^\w\d.]{1,5}', discipline_raw)
                discipline_code = discipline_parts[0]
                if len(discipline_parts) < 2:
                    discipline = ''
                else:
                    discipline = discipline_parts[1]

            if profile_raw != str():
                profile = remove_keywords(profile_raw, key_words)
                study_raw = remove_keywords(study_raw, key_words).strip()

                group_regex_fnd = re.search('(\d{1,2}\.?){3}', study_raw)

                group_code_span = (0,0) if group_regex_fnd is None else group_regex_fnd.span()
                group_code = study_raw[group_code_span[0]:group_code_span[1]]
                study_i = study_raw[group_code_span[1]:].strip()
                study = (group_code, study_i)
            if variant_raw != str():
                variant_span = re.search('(?<=вариант) *\d{1,}', variant_raw, re.IGNORECASE)
                variant_span = (0, len(variant_raw)) if variant_span is None else variant_span.span()
                variant = variant_raw[variant_span[0]:variant_span[1]].strip()

            if qual_raw != str():
                qual = remove_keywords(qual_raw, key_words)
                qual_span = re.search('\w+', qual).span()
                qual = qual[qual_span[0]:qual_span[1]]

            doc_info = Work(student=student_info_raw,
                            number=number, topic=topic, faculty=faculty,
                            chair=chair, group=(group, study[0]), year=year,
                            files=documents, study=study[1], profile=profile,
                            discipline=(discipline, discipline_code), variant=variant,
                            qualification=qual)
            doc_info.lines['student'] = get_line(student_info_raw, first_page_text)
            doc_info.lines['faculty'] = get_line(faculty, first_page_text)
            doc_info.lines['chair'] = get_line(chair_raw, first_page_text)
            doc_info.lines['topic'] = get_line('на тему:', first_page_text)
            doc_info.lines['group'] = get_line(group, first_page_text)
            doc_info.lines['topic'] = doc_info.lines['topic'] + 1 if doc_info.lines['topic'] != -1 else -1
            doc_info.lines['year'] = get_line(year_info, first_page_text)

            parsed_info.append(doc_info)
    print('That array of files was read')
    return parsed_info


def get_document_first_page(document):
    first_page_text = []
    for p in document.paragraphs[:50]:
        if len(p.text) > 0:
            first_page_text.append(p.text)
    return first_page_text


def define_line_links(doc_info: Work, search_source: List[str]=list()):
    doc_info.lines['student'] = get_line(doc_info['student'], search_source)
    doc_info.lines['faculty'] = get_line(doc_info['faculty'], search_source)
    doc_info.lines['chair'] = get_line(doc_info['chair'], search_source)
    doc_info.lines['topic'] = get_line('на тему:', search_source)
    doc_info.lines['group'] = get_line(doc_info['group'], search_source)
    doc_info.lines['topic'] = doc_info.lines['topic'] + 1 if doc_info.lines['topic'] != -1 else -1
    doc_info.lines['year'] = get_line(doc_info['year'], search_source)


def get_common_info(works: List[Work]):
    """
    Собирает общую информацию о работах
    :param works: работы
    :return: словарь, сключи которого - типы общей информации (тема, факультет, группа)
    """
    def create_group(input_list: List[Work], key: str):
        if input_list is None or len(input_list) == 0:
            return list()
        new_group = groupby(sorted(input_list, key=itemgetter(key)), key=itemgetter(key))
        new_group = [(name, list(elements)) for name, elements in new_group]
        new_group = sorted(new_group, key=lambda x: len(x[1]), reverse=True)
        return new_group

    teacher_group = create_group(works, 'teacher')
    topic_group = create_group(works, 'topic')
    faculty_group = create_group(works, 'faculty')
    chair_group = create_group(works, 'chair')
    gr_group = create_group(works, 'group')
    gr_code_group = create_group(works, 'group_code')
    year_group = create_group(works, 'year')
    work_type_group = create_group(works, 'work_type')

    qual_group = create_group(works, 'qual')
    discipline_group = create_group(works, 'discipline')
    discipline_code_group = create_group(works, 'discipline_code')
    profile_group = create_group(works, 'profile')
    study_group = create_group(works, 'study')
    teacher_duty_group = create_group(works, 'teacher_duty')
    return {
        'teacher': teacher_group,
        'topic': topic_group,
        'faculty': faculty_group,
        'chair': chair_group,
        'group': gr_group,
        'group_code': gr_code_group,
        'year': year_group,
        'work_type': work_type_group,

        'qual': qual_group,
        'discipline': discipline_group,
        'discipline_code': discipline_code_group,
        'profile': profile_group,
        'study': study_group,
        'teacher_duty': teacher_duty_group
    }


def save_changes(works: List[Work], common_info: Dict[str, str], dir_path: str, change_title_pages: bool = False, call_dialog_method = lambda path, student: QMessageBox.Yes):
    """
    Создание титульных страниц
    :param works: работы
    :param common_info: общая информация о работах
    :param dir_path: исходная директория
    :param change_title_pages: флаг перезаписи титульников
    :param call_dialog_method: метод для отображения MessageBox с вопросом о перезаписи файлов с именами-дубликатами
    :return:
    """
    print('Принято решение {}изменять титульники'.format('' if change_title_pages else 'не '))

    folder_name = '{faculty}_{chair_short}_{group}_{work_type}'.format(**common_info)
    folder_path = os.path.join(Path(dir_path).parent, 'Результат', folder_name)
    if not os.path.exists(folder_path):
        os.makedirs(folder_path, exist_ok=True)

    for work in works:
        document_title, title_page_document = make_title_page(work, common_info)
        if isinstance(work.files, str):
            raise ValueError('Вместо списка здесь строка')
        else:
            existing_files = [os.path.join(dir_path, file) for file in work.files]
            report_file = [file for file in existing_files if file.endswith('docx')]
            report_file = next(iter(report_file))

        try:
            report_document = Document(report_file)
        except:
            continue
        new_report_document = report_document
        if all([line == 0 for line in work.lines.values()]):
            first_page_contents = [par.text for par in report_document.paragraphs[:50]]
            define_line_links(work, first_page_contents)
        if work.has_changed and change_title_pages:
            old_title_page_end_pos = max(work.lines.values()) if max(work.lines.values()) > 1 else 0  # work.lines['year']
            new_report_document = insert_title_page(title_page_document=title_page_document,
                                                    main_document=report_document,
                                                    replace_start_point=old_title_page_end_pos)
            # font = new_report_document.styles['Normal'].font
            # font.name = 'Times New Roman'
        report_document_new_path = os.path.join(folder_path, document_title)
        if os.path.exists(report_document_new_path):
            rewrite_choice = call_dialog_method(document_title, work['student'])
            if rewrite_choice == QMessageBox.Abort:
                print('Найден файл с именем-дубликатом, отмена всего')
                break
            elif rewrite_choice == QMessageBox.Yes:
                print('Найден файл с именем-дубликатом, перезаписываем')
                os.remove(report_document_new_path)
            else:
                print('Найден файл с именем-дубликатом, не трогаем')
                continue
        new_report_document.save(report_document_new_path)
        renamed_file_paths = [document_title]

        if len(existing_files) > 1:
            other_files = list(list(filter(lambda x: not x.endswith('docx'), existing_files)))
            number = 2
            for file in other_files:
                ext = os.path.splitext(file)[-1]
                new_name = '{0}_{1:2}{2}'.format(document_title[:-5], number, ext)  # '{0}_{1:02}{2}'
                shutil.copy2(file, os.path.join(folder_path, new_name))
                renamed_file_paths.append(new_name)
                number += 1
        work.files = renamed_file_paths
