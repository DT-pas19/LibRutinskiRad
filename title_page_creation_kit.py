from typing import Dict, Tuple

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from document_structure import Work
from helpers import quote


__university__ = 'Саратовский государственный технический университет имени Гагарина Ю.А.'
__work_type_desc__ = {
    1: 'Выпускная квалификационная работа',
    2: 'Дипломная работа',
    3: 'Дипломный проект',
    4: 'Магистерская диссертация',
    5: 'Отчет по практике',
    6: 'Курсовой проект',
    7: 'Контрольная работа',
    8: 'Расчетно-графическая работа',
    9: 'Курсовая работа',
    10: 'Научная квалификационная работа',
    11: 'Научный доклад'
}


def make_title_page(work: Work, common_info: Dict[str, str]) -> Tuple[str, Document]:
    document = Document()
    caption_paragraphs = list()
    topic_paragraphs = list()
    persons_paragraph = list()
    caption_paragraphs.append(document.add_paragraph('Министерство образования и науки Российской Федерации'))
    caption_paragraphs.append(document.add_paragraph('Федеральное государственное бюджетное учреждение'))
    caption_paragraphs.append(document.add_paragraph('высшего профессионального образования'))
    caption_paragraphs.append(document.add_paragraph(quote(__university__)))
    caption_paragraphs.append(document.add_paragraph())

    caption_paragraphs.append(document.add_paragraph('Факультет: {}'.format(common_info['faculty'])))
    caption_paragraphs.append(document.add_paragraph('Кафедра {}'.format(quote(common_info['chair']))))

    document.add_paragraph()
    document.add_paragraph()

    type_desc = __work_type_desc__[int(work['work_type'])]
    topic_paragraphs.append(document.add_paragraph(type_desc))

    if work['work_type'] == 7:
        topic_paragraphs.append(document.add_paragraph('на тему:'))
        topic_paragraphs.append(document.add_paragraph(quote(work['topic'])))
    if work['work_type'] == 8:
        discipline = work['discipline_code'] + ' ' + quote(work['discipline'])
        # 'Б.1.1.7 ' + quote('Информатика')
        direction = work['group_code'] + quote(work['study'])
        dir_profile = work['profile']
        qualification = work['qual']
        var = work['variant'] if 'variant' in work.keys() else 0
        topic_paragraphs.append(document.add_paragraph('по дисциплине {}'.format(discipline)))
        topic_paragraphs.append(document.add_paragraph('направление подготовки {}'.format(direction)))
        topic_paragraphs.append(document.add_paragraph('профиль {}'.format(dir_profile)))
        topic_paragraphs.append(document.add_paragraph('Квалификация (степень) {}'.format(qualification)))
        topic_paragraphs.append(document.add_paragraph('Тема {} Вариант {}'.format(quote(work['topic']), var)))

    document.add_paragraph()
    document.add_paragraph()
    document.add_paragraph()
    document.add_paragraph()

    persons_paragraph.append(document.add_paragraph('Выполнил:'))
    persons_paragraph.append(document.add_paragraph('студент группы {}'.format(work['group'])))
    persons_paragraph.append(document.add_paragraph(work['student']))
    persons_paragraph.append(document.add_paragraph('Зачет.кн. №{}'.format(work['number'])))
    persons_paragraph.append(document.add_paragraph('Проверил:'))

    persons_paragraph.append(document.add_paragraph(common_info['teacher_duty']))
    persons_paragraph.append(document.add_paragraph(common_info['teacher']))
    document.add_paragraph()
    document.add_paragraph()
    year_par = document.add_paragraph('Саратов {}'.format(common_info['year']))
    document.add_page_break()

    # font = document.styles['Normal'].font
    # font.name = 'Times New Roman'
    # font.size = Pt(14)
    # <editor-fold desc="Alignment settings">
    for par in caption_paragraphs + topic_paragraphs:
        par.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for par in persons_paragraph:
        par.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    year_par.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # </editor-fold>

    document_name = '{number}_{group}_{year}_{work_type}.docx'.format(**work)
    return document_name, document