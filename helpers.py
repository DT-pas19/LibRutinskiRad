import re
import docx
from typing import Dict
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.shared import Pt, Inches


def quote(string: str):
    if len(string) == 0:
        return str()
    result = ''.join(('\u00ab', string, '\u00bb'))
    return result


def remove_keywords(text: str, key_words: Dict[str, str]) -> str:
    """
    Удаление ключевых слов + прочих символов (:()»«"\' \t,) из строки
    :param key_words: ключевые слова
    :param text: строка
    :return: новая строка
    """
    if len(text) == 0:
        return str()
    text = replace_whitespace(text)
    # key_list = [key for key in key_words if key.lower() in text.lower()]
    key_list = [key for key in key_words.values() if isinstance(key, str) and key.lower() in text.lower()]
    key_list += [key for key in key_words.values() if isinstance(key, tuple) and any(k.lower() in text.lower() for k in key)]
    if any([isinstance(key, tuple) for key in key_list]):
        key_list_r = list()
        for key in key_list:
            if isinstance(key, str):
                key_list_r.append(key)
            elif isinstance(key, tuple):
                key_list_r += list(key)
        key_list = key_list_r
    reduced_text = text
    split_result = re.split('\w*{0}\w*'.format('\w*|\w*'.join(key_list)), reduced_text, flags=re.IGNORECASE)
    reduced_text = ' '.join(split_result).strip(':()»«"\' \t,')
    reduced_text = ' '.join([sub for sub in re.split('[:()»«]', reduced_text) if len(sub) > 0])
    if '\n' in reduced_text:
        reduced_text.replace('\n', '')
    return reduced_text


def get_bracket_less_line(line: str):
    """
    Возвращает строку без опоясывающих скобок
    :param line: строка
    :return: строка без скобок
    """
    if '(' in line and ')' in line:
        return line[1:-1]
    return line


def replace_whitespace(line: str):
    """
    Возвращает строку без лишних пробельных символов
    :param line: строка
    :return: строка без лишних пробелов
    """
    if not isinstance(line, str) or len(line) == 0:
        return line
    parts = line.strip().split()
    recombined = ' '.join(parts)
    return recombined


def insert_title_page(title_page_document: Document, main_document: Document, replace_start_point: int=0):
    """
    Вставка титульного листа в документ
    :param title_page_document: сформированные параграфы титульного листа
    :param main_document: главный документ
    :param replace_start_point: точка, до которой следует заменять текст
    :return:
    """

    if not isinstance(replace_start_point, int):
        return
    replace_start_point = 0 if replace_start_point < 0 else replace_start_point
    modified_main_document = main_document

    try:
        style = modified_main_document.styles.add_style('TitlePage', WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        style = modified_main_document.styles['TitlePage']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(16)
    paragraph_format = style.paragraph_format
    paragraph_format.left_indent = Inches(0)
    paragraph_format.first_line_indent = Inches(0)
    paragraph_format.space_before = Pt(16)
    paragraph_format.widow_control = True

    counter = 0
    insert_point = modified_main_document.paragraphs[replace_start_point]

    for par in title_page_document.paragraphs:
        if counter <= replace_start_point:
            modified_main_document.paragraphs[counter].text = par.text
            modified_main_document.paragraphs[counter].style = modified_main_document.styles['TitlePage']
            modified_main_document.paragraphs[counter].paragraph_format.alignment = par.paragraph_format.alignment
            counter += 1
        else:
            inserted_par = insert_point.insert_paragraph_before(text=par.text, style='TitlePage')
            inserted_par.paragraph_format.alignment = par.paragraph_format.alignment
    for i in range(replace_start_point - counter + 1):
        modified_main_document.paragraphs[counter].text = str()
        counter += 1

    insert_point.insert_paragraph_before(text='\n', style='TitlePage')

    # insert_point = modified_main_document.paragraphs[0]
    # insert_point.insert_paragraph_before(text='', style='Normal')
    # modified_main_document.paragraph_format.alignment = par.paragraph_format.alignment
    return modified_main_document
